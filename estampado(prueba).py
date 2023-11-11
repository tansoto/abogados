import sys
from PyQt6.QtWidgets import QApplication, QMainWindow, QVBoxLayout, QWidget, QPushButton, QTextEdit, QFileDialog
from docx import Document
from reportlab.pdfgen import canvas

class EstampadoApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Editor de Word')
        self.setGeometry(100, 100, 600, 400)

        self.central_widget = QWidget(self)
        self.setCentralWidget(self.central_widget)

        self.layout = QVBoxLayout()

        self.text_edit = QTextEdit(self)
        self.layout.addWidget(self.text_edit)

        # Botones
        self.btn_open = QPushButton('Abrir Documento', self)
        self.btn_open.clicked.connect(self.open_document)
        self.layout.addWidget(self.btn_open)

        self.btn_save = QPushButton('Guardar Documento como word', self)
        self.btn_save.clicked.connect(self.save_document)
        self.layout.addWidget(self.btn_save)

        self.btn_save_pdf = QPushButton('Guardar como PDF', self)
        self.btn_save_pdf.clicked.connect(self.save_as_pdf)
        self.layout.addWidget(self.btn_save_pdf)

        self.btn_back = QPushButton('Volver', self)
        self.btn_back.clicked.connect(self.close)
        self.layout.addWidget(self.btn_back)

        self.central_widget.setLayout(self.layout)

        self.current_document_path = None

    def open_document(self):
        # Abre un cuadro de diálogo para seleccionar un documento Word
        file_path, _ = QFileDialog.getOpenFileName(self, "Abrir Documento Word", "", "Documentos Word (*.docx)")

        if file_path:
            # Almacena la ruta del documento actual
            self.current_document_path = file_path

            # Lee el contenido del documento Word y lo muestra en el QTextEdit
            document = Document(file_path)
            content = ""
            for paragraph in document.paragraphs:
                content += paragraph.text + '\n'
            self.text_edit.setPlainText(content)

    def save_document(self):
        if self.current_document_path:
            # Guarda el contenido actual en un documento Word
            document = Document()
            for line in self.text_edit.toPlainText().split('\n'):
                document.add_paragraph(line)
            document.save(self.current_document_path)

            # Muestra un mensaje de éxito en la barra de estado
            self.statusBar().showMessage("Documento guardado con éxito.")
        else:
            # Si no hay una ruta de documento actual, solicita una nueva ubicación
            self.save_as_document()

    def save_as_document(self):
        # Abre un cuadro de diálogo para seleccionar la ubicación y el nombre del nuevo documento Word
        file_path, _ = QFileDialog.getSaveFileName(self, "Guardar Documento Word", "", "Documentos Word (*.docx)")

        if file_path:
            # Almacena la nueva ruta del documento
            self.current_document_path = file_path
            # Guarda el documento
            self.save_document()

    def save_as_pdf(self):
        if self.current_document_path:
            # Abre un cuadro de diálogo para seleccionar la ubicación y el nombre del nuevo documento PDF
            pdf_file_path, _ = QFileDialog.getSaveFileName(self, "Guardar como PDF", "", "Archivos PDF (*.pdf)")

            if pdf_file_path:
                # Crea un nuevo documento PDF y agrega el contenido del documento Word
                document = Document(self.current_document_path)
                pdf_canvas = canvas.Canvas(pdf_file_path)
                pdf_canvas.setFont("Helvetica", 12)

                for paragraph in document.paragraphs:
                    pdf_canvas.drawString(50, pdf_canvas._pagesize[1] - 50, paragraph.text)
                    pdf_canvas.showPage()

                pdf_canvas.save()

                # Muestra un mensaje de éxito en la barra de estado
                self.statusBar().showMessage("Documento guardado como PDF con éxito.")
        else:
            # Si no hay una ruta de documento actual, muestra un mensaje en la barra de estado
            self.statusBar().showMessage("Abre o guarda un documento Word antes de convertirlo a PDF.")

def main():
    app = QApplication(sys.argv)
    window = EstampadoApp()
    window.show()
    sys.exit(app.exec())

if __name__ == '__main__':
    main()
